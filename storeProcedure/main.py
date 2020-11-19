import json
import os
import re

fileInfo = json.load(open("./fileInfo.json", "r", encoding="utf-8"))

headerMapDict = {
    '說明':'description',
    '輸入參數':'inputParameter',
    '輸出結果':'outputResult',
    '錯誤':'errorMsg',
}
inputParameterMapDict = {
    '名稱':'inputName',
    '型態':'inputType',
}
errorMsgMapDict = {
    '代碼':'errno',
    '訊息':'errMsg'
}
headerList = [key for key, _ in headerMapDict.items()]
bodyDict = {value:headerList.index(key) for key, value in headerMapDict.items()}

#解析sql裡的SP
with open(fileInfo["sqlFileName"], "r", encoding = "utf-8") as inputContent:
        lines = inputContent.readlines() #讀進全部
        inTarget = False #是否在Procedure的範圍
        target = {}
        for line in lines: 
            lineU = line.upper() # 都轉成大寫，統一比對
            lineUSplit = lineU.split()
            #是否即將脫離Procedure的範圍
            if "$$" in line: 
                inTarget = False
            #若是在表中，代表現在這行為欄位資訊
            elif inTarget:
                x={} 
                #抓出Error訊息
                if "MYSQL_ERRNO" in lineU:
                    x['errno'] = lineUSplit[lineUSplit.index("MYSQL_ERRNO") + 2].strip(",")
                    x['errMsg'] = ' '.join(line.split()[lineUSplit.index("MESSAGE_TEXT") + 2:]).strip("',;") if "MESSAGE_TEXT" in lineU else None
                    target[targetName]['errorMsg'].append(x)
                elif "CALL" in lineU:
                    target[targetName]['call'].append(line.strip().split('`')[1])
            elif "PROCEDURE" in lineU:
                targetName = line.split()[3].strip("`") #前後拿掉`符號
                target[targetName] = {}
                target[targetName]['errorMsg'] = []
                target[targetName]['inputParameter'] = []
                target[targetName]['call'] = []
                target[targetName]['outputResult'] = ""
                target[targetName]['description'] = ""
                indices = [i for i, element in enumerate(lineUSplit) if element == "(IN" or element == "IN"]
                #抓出輸入參數
                for index in indices:
                    x = {}
                    x['inputName'] = line.split()[index + 1].strip('`,)')
                    x['inputType'] = line.split()[index + 2].strip('`,)')
                    target[targetName]['inputParameter'].append(x)
                inTarget = True
                
#若在SP中有call其他SP，要將其他SP中的錯誤納進這個SP
for targetName, targetProp in target.items():
    for calledSp in targetProp['call']:
        targetProp['errorMsg'] += target[calledSp]['errorMsg']
    #排序，號碼小的在前面
    t = {element['errno']:element['errMsg'] for element in sorted(targetProp['errorMsg'], key = lambda x: (x['errno'], x['errMsg']))}
    targetProp['errorMsg'] = [{'errno':key, 'errMsg': value} for key, value in t.items()]



#輸出成docx
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.oxml.ns import qn
from docx.oxml.shared import OxmlElement
from docx.shared import Cm, Inches
from docx.oxml.ns import nsdecls
from docx.oxml import parse_xml

document = Document()

if os.path.isfile(fileInfo["descriptionFileName"]):
    targetDes = json.load(open(fileInfo["descriptionFileName"], "r", encoding="utf-8"))
else:
    targetDes = {}

for targetName, targetProp in target.items():
    if targetName in targetDes:
        targetProp['outputResult'] = targetDes[targetName]['outputResult']
        targetProp['description'] = targetDes[targetName]['description']
        #加入輸入參數的註解
        for inputObj in targetProp['inputParameter']:
            if inputObj['inputName'] in targetDes[targetName]['inputParameter']:
                inputObj['description'] = targetDes[targetName]['inputParameter'][inputObj['inputName']]
            
for targetName, targetProp in target.items():
    document.add_heading(targetName, level=1)
    
    tb = document.add_table(rows=len(headerList), cols=2, style='Table Grid')
      
    #增加固定欄寬節點
    child = OxmlElement('w:tblLayout')
    tb._element.tblPr.append(child)
    tb._element.tblPr.tblLayout.set(qn('w:type'), 'fixed')
    
    widths = (Cm(2.2), Cm(13))
    for row in tb.rows:
        for idx, width in enumerate(widths):
            row.cells[idx].width = width
            
    hdrCells = tb.columns[0].cells
    #表格的頭部
    for i in range(0, len(headerList)):
        #加上背景色
        shading_elm = parse_xml(r'<w:shd {} w:fill="DBDBDB"/>'.format(nsdecls('w')))
        hdrCells[i]._tc.get_or_add_tcPr().append(shading_elm)
        p = hdrCells[i].paragraphs[0] #拿原本的paragraphs來操作
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT #置左
        run = p.add_run(headerList[i]) #追加字串
        run.bold = True #粗體
        run.font.size = Pt(12)  #字的大小
        run.font.name = 'New Times Roman' #英文字型
        run._element.rPr.rFonts.set(qn('w:eastAsia'), u'標楷體') #中文字型      
    
    columnCells = tb.columns[1].cells  
    #表格的身體
    for propName, propValue in targetProp.items():
        if propName not in bodyDict:
            continue
        p = columnCells[bodyDict[propName]].paragraphs[0]
        if propValue is None:
            continue
        elif type(propValue) is str:
            run = p.add_run(propValue)
        elif propName == headerMapDict['輸入參數']:
            for element in propValue:
                run = p.add_run(element['inputName'])
                run.bold = True
                run = p.add_run("："+ element['inputType'] + ("，" + element['description'] if element['description'] != "" else "。" ))
                run.font.size = Pt(10)  #字的大小
                run.font.name = 'New Times Roman' #英文字型
                run._element.rPr.rFonts.set(qn('w:eastAsia'), u'標楷體') #中文字型
                if propValue.index(element) != (len(propValue) - 1):
                    run.add_break(WD_BREAK.LINE)
        elif propName == headerMapDict['錯誤']:
            if len(propValue) == 0:
                #就算沒資料，也要add_run，不然標題的字型大小會異常
                run = p.add_run()
                run.font.size = Pt(10)  #字的大小
                continue
            for element in propValue:
                run = p.add_run("#" + element['errno'] + " - " + element['errMsg'])
                run.font.size = Pt(10)  #字的大小
                run.font.name = 'New Times Roman' #英文字型
                run._element.rPr.rFonts.set(qn('w:eastAsia'), u'標楷體') #中文字型
                if propValue.index(element) != (len(propValue) - 1):
                    run.add_break(WD_BREAK.LINE)
                
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        run.font.size = Pt(10)  #字的大小
        run.font.name = 'New Times Roman' #英文字型
        run._element.rPr.rFonts.set(qn('w:eastAsia'), u'標楷體') #中文字型


#取出所有錯誤代碼與訊息
errorAll = []
for targetName, targetProp in target.items():
    errorAll += targetProp['errorMsg']

errs = {}
for x in errorAll:
    #使用set過濾同代碼、同訊息
    if x['errno'] not in errs:
        errs[x['errno']] = {x['errMsg']}
    else:
        errs[x['errno']].add(x['errMsg'])
#依照errno排序
errorAll = {errno:errs[errno] for errno in sorted(errs, key = lambda k: (k))}

headerList = ["錯誤代碼", "錯誤說明"]
document.add_heading("錯誤對照表", level=1)
errorClassNo = int(sorted(errorAll.keys())[-1]) // 100 - 100 
tb = document.add_table(rows=len(errorAll) + errorClassNo + 1, cols=2, style='Table Grid')
widths = (Cm(2.2), Cm(13))
for row in tb.rows:
    for idx, width in enumerate(widths):
        row.cells[idx].width = width
hdrCells = tb.rows[0].cells
#表格的頭部
for i in range(0, len(headerList)):
    p = hdrCells[i].paragraphs[0] #拿原本的paragraphs來操作
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT #置左
    run = p.add_run(headerList[i]) #追加字串
    run.bold = True #粗體
    run.font.size = Pt(12)  #字的大小
    run.font.name = 'New Times Roman' #英文字型
    run._element.rPr.rFonts.set(qn('w:eastAsia'), u'標楷體') #中文字型
errorClass = {
    1:"輸入參數錯誤",
    2:"新增帳號錯誤",
    3:"初始化錯誤",
    4:"帳號權限錯誤",
}

i = 0
prevClass = 0
for errno, errMsg in errorAll.items():
    i += 1
    currClass = int(errno) // 100 - 100 #第幾個類別
    #加入錯誤類別名稱列
    if currClass != prevClass:
        tb.rows[i].cells[0].merge(tb.rows[i].cells[1]) #合併欄位
        run = tb.rows[i].cells[0].paragraphs[0].add_run(errorClass[currClass]) 
        #加上背景色
        shading_elm = parse_xml(r'<w:shd {} w:fill="FFC000"/>'.format(nsdecls('w')))
        tb.rows[i].cells[0]._tc.get_or_add_tcPr().append(shading_elm)
        run.bold = True #粗體
        run.font.size = Pt(10)  #字的大小
        run.font.name = 'New Times Roman' #英文字型
        run._element.rPr.rFonts.set(qn('w:eastAsia'), u'標楷體') #中文字型
        i += 1
    #每個錯誤項
    p = tb.rows[i].cells[0].paragraphs[0]
    run = p.add_run(errno)
    run.font.size = Pt(10)  #字的大小
    run.font.name = 'New Times Roman' #英文字型
    run._element.rPr.rFonts.set(qn('w:eastAsia'), u'標楷體') #中文字型
    errMsg = list(errMsg)
    p = tb.rows[i].cells[1].paragraphs[0]
    for msg in errMsg:
        run = p.add_run(msg)
        run.font.size = Pt(10)  #字的大小
        run.font.name = 'New Times Roman' #英文字型
        run._element.rPr.rFonts.set(qn('w:eastAsia'), u'標楷體') #中文字型
        if errMsg.index(msg) != (len(errMsg) - 1):
            run.add_break(WD_BREAK.LINE)
            
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run.font.size = Pt(10)  #字的大小
    run.font.name = 'New Times Roman' #英文字型
    run._element.rPr.rFonts.set(qn('w:eastAsia'), u'標楷體') #中文字型
    prevClass = currClass
document.save(fileInfo['docxFileName'])


#docx 轉 pdf

from docx2pdf import convert
convert(fileInfo['docxFileName'])
os.rename(fileInfo['docxFileName'].split('.')[0] + '.pdf', fileInfo['pdfFileName'])