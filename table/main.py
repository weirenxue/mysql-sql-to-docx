import json
import os

fileInfo = json.load(open("./fileInfo.json", "r", encoding="utf-8"))

mapDict = {
    '欄位名稱':'colName',
    '資料型態':'type',
    'Nullable':'nullable',
    '默認值':'default',
    '欄位說明':'comment'
}
headerList = [key for key, _ in mapDict.items()]
columnDict = {value:headerList.index(key) for key, value in mapDict.items()}

#解析sql裡的Table，sql中可以有TABLE DDL，不可以有view、sp等
with open(fileInfo["sqlFileName"], "r", encoding = "utf-8") as inputContent:
        lines = inputContent.readlines() #讀進全部
        inTable = False #是否在CREATE TABLE的範圍
        table = {}
        for line in lines: 
            #是否即將脫離CREATE TABLE的範圍
            if ";" in line: 
                inTable = False
            #若是在表中，代表現在這行為欄位資訊
            elif inTable:
                x={}
                lineU = line.upper() # 都轉成大寫，統一比對
                col = lineU.split()
                x[mapDict['欄位名稱']] = col[0].strip("`")
                x[mapDict['資料型態']] = col[1] + (" UNSIGNED" if "UNSIGNED" in lineU else "")
                x[mapDict['Nullable']] = False if "NOT NULL" in lineU else True
                x[mapDict['默認值']] = col[col.index("DEFAULT") + 1] if "DEFAULT" in lineU else None
                x[mapDict['欄位說明']] = col[-1].strip("',") if "COMMENT" in lineU else False
                #x['collate'] = col[col.index("DEFAULT") + 1] if "COLLATE" in lineU else None
                table[tableName].append(x)

            elif "CREATE TABLE" in line.upper():
                tableName = line.split()[2].strip("`") #前後拿掉`符號
                table[tableName] = []
                inTable = True
            
            
                
#輸出成docx
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml.shared import OxmlElement

document = Document()

if os.path.isfile(fileInfo["tableDescriptionFileName"]):
    tableDes = json.load(open(fileInfo["tableDescriptionFileName"], "r", encoding="utf-8"))
else:
    tableDes = {}
    
for tableName, columns in table.items():
    document.add_heading(tableName, level=1)
    
    if tableName in tableDes:
        p = document.add_paragraph(tableDes[tableName])
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    
    tb = document.add_table(rows=1 + len(columns), cols=len(headerList), style='Table Grid')
      
    #增加固定欄寬節點
    child = OxmlElement('w:tblLayout')
    tb._element.tblPr.append(child)
    tb._element.tblPr.tblLayout.set(qn('w:type'), 'fixed')
              
    hdr_cells = tb.rows[0].cells
    #表格的頭部
    for i in range(0, len(headerList)):
        p = hdr_cells[i].paragraphs[0] #拿原本的paragraphs來操作
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER #置中
        run = p.add_run(headerList[i]) #追加字串
        run.bold = True #粗體
        run.font.size = Pt(12)  #字的大小
        run.font.name = 'New Times Roman' #英文字型
        run._element.rPr.rFonts.set(qn('w:eastAsia'), u'標楷體') #中文字型
    #表格的身體
    for column in columns:
        row_cells = tb.rows[columns.index(column) + 1].cells
        for key, index in columnDict.items():
            p = row_cells[index].paragraphs[0]
            value = column[key]
            if value is None:
                continue
            elif type(value) is bool:
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                run = p.add_run("v" if column[key] else "")
            else:
                run = p.add_run(column[key])
            run.font.size = Pt(10)  #字的大小
            run.font.name = 'New Times Roman' #英文字型
            run._element.rPr.rFonts.set(qn('w:eastAsia'), u'標楷體') #中文字型

document.save(fileInfo['docxFileName'])


#docx 轉 pdf
from docx2pdf import convert
convert(fileInfo['docxFileName'])
os.rename(fileInfo['docxFileName'].split('.')[0] + '.pdf', fileInfo['pdfFileName'])